from pathlib import Path
from typing import Tuple

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


FONT_NAME = "Arial"
FONT_SIZE_PT = 10
LINE_SPACING = 1
SOURCE_DIR = Path("Templates_markdown")
TARGET_DIR = Path("Templates_docx")


def configure_base_style(document: Document) -> None:
    """Configure the default paragraph style to match the requested formatting."""
    style = document.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(FONT_SIZE_PT)

    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = LINE_SPACING


def append_run(
    paragraph,
    text: str,
    bold: bool,
    italic: bool,
    force_italic: bool,
    font_size_pt: int,
) -> None:
    """Add a run to the paragraph with the correct font settings."""
    if not text:
        return
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic or force_italic
    run.font.name = FONT_NAME
    run.font.size = Pt(font_size_pt)


def normalize_heading(line: str) -> Tuple[str, bool]:
    """Return the line text and whether it should be treated as a heading (bold)."""
    stripped = line.lstrip()
    if stripped.startswith("#"):
        return stripped.lstrip("#").strip(), True
    return line, False


def add_markdown_paragraph(
    document: Document,
    raw_line: str,
    *,
    alignment: WD_ALIGN_PARAGRAPH,
    force_italic: bool = False,
    font_size_pt: int = FONT_SIZE_PT,
) -> None:
    """Convert a single markdown line into a DOCX paragraph."""
    text, heading = normalize_heading(raw_line)
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = alignment
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = LINE_SPACING

    if text == "":
        paragraph.add_run("")
        return

    buffer = []
    bold = heading
    italic = False
    i = 0
    while i < len(text):
        if text.startswith("**", i):
            append_run(
                paragraph, "".join(buffer), bold, italic, force_italic, font_size_pt
            )
            buffer = []
            bold = not bold
            i += 2
            continue
        if text.startswith("__", i):
            append_run(
                paragraph, "".join(buffer), bold, italic, force_italic, font_size_pt
            )
            buffer = []
            bold = not bold
            i += 2
            continue
        if text[i] == "*":
            append_run(
                paragraph, "".join(buffer), bold, italic, force_italic, font_size_pt
            )
            buffer = []
            italic = not italic
            i += 1
            continue
        if text[i] == "_":
            append_run(
                paragraph, "".join(buffer), bold, italic, force_italic, font_size_pt
            )
            buffer = []
            italic = not italic
            i += 1
            continue

        buffer.append(text[i])
        i += 1

    append_run(
        paragraph, "".join(buffer), bold, italic, force_italic, font_size_pt
    )


def convert_file(md_path: Path, output_path: Path) -> None:
    """Convert a markdown file into a DOCX file."""
    document = Document()
    configure_base_style(document)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    if not lines:
        document.add_paragraph("", style="Normal")
    else:
        first_written = next((i for i, line in enumerate(lines) if line.strip()), None)
        last_written = (
            len(lines) - 1
            - next((i for i, line in enumerate(reversed(lines)) if line.strip()), None)
            if any(line.strip() for line in lines)
            else None
        )

        for idx, line in enumerate(lines):
            alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            force_italic = False
            font_size_pt = FONT_SIZE_PT
            if first_written is not None and idx == first_written:
                alignment = WD_ALIGN_PARAGRAPH.CENTER
            if last_written is not None and idx == last_written:
                alignment = WD_ALIGN_PARAGRAPH.CENTER
                force_italic = True
                font_size_pt = 8

            add_markdown_paragraph(
                document,
                line,
                alignment=alignment,
                force_italic=force_italic,
                font_size_pt=font_size_pt,
            )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    if not SOURCE_DIR.exists():
        raise FileNotFoundError(f"Source folder not found: {SOURCE_DIR}")

    TARGET_DIR.mkdir(parents=True, exist_ok=True)
    for md_file in sorted(SOURCE_DIR.glob("*.md")):
        output_file = TARGET_DIR / f"{md_file.stem}.docx"
        convert_file(md_file, output_file)


if __name__ == "__main__":
    main()
