#!/usr/bin/env python3
"""
Script para converter arquivos .docx e .rtf para markdown preservando formatação.
"""

import os
import sys
from pathlib import Path
import re

try:
    from docx import Document
except ImportError:
    print("Instalando python-docx...")
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document

try:
    from striprtf.striprtf import rtf_to_text
except ImportError:
    print("Instalando striprtf...")
    os.system(f"{sys.executable} -m pip install striprtf")
    from striprtf.striprtf import rtf_to_text


def convert_docx_to_markdown(docx_path):
    """Converte um arquivo .docx para markdown preservando formatação."""
    try:
        doc = Document(docx_path)
        markdown_lines = []
        
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                markdown_lines.append("")
                continue
            
            # Processa runs para preservar formatação
            text_parts = []
            for run in paragraph.runs:
                text = run.text
                if not text:
                    continue

                # Aplica formatação em Markdown puro (sem HTML)
                if run.bold:
                    text = f"**{text}**"
                if run.italic:
                    text = f"*{text}*"
                if run.underline:
                    text = f"__{text}__"

                text_parts.append(text)
            
            # Se não há runs formatados, usa o texto do parágrafo
            if not text_parts:
                para_text = paragraph.text
            else:
                para_text = "".join(text_parts)
            
            # Verifica estilo de parágrafo
            style = paragraph.style.name.lower()
            
            # Títulos
            if 'heading' in style or 'title' in style:
                level = 1
                if 'heading 1' in style or 'title' in style:
                    level = 1
                elif 'heading 2' in style:
                    level = 2
                elif 'heading 3' in style:
                    level = 3
                elif 'heading 4' in style:
                    level = 4
                elif 'heading 5' in style:
                    level = 5
                elif 'heading 6' in style:
                    level = 6
                
                markdown_lines.append(f"{'#' * level} {para_text}")
            else:
                markdown_lines.append(para_text)
        
        # Processa tabelas
        for table in doc.tables:
            markdown_lines.append("")
            # Cabeçalho
            header_row = table.rows[0]
            header_cells = [cell.text.strip() for cell in header_row.cells]
            markdown_lines.append("| " + " | ".join(header_cells) + " |")
            markdown_lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
            
            # Linhas de dados
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                markdown_lines.append("| " + " | ".join(cells) + " |")
            markdown_lines.append("")
        
        return "\n".join(markdown_lines)
    
    except Exception as e:
        return f"# Erro ao converter {docx_path}\n\nErro: {str(e)}"


def convert_rtf_to_markdown(rtf_path):
    """Converte um arquivo .rtf para markdown preservando formatação básica."""
    try:
        with open(rtf_path, 'rb') as f:
            rtf_content = f.read()
        
        # Tenta diferentes encodings
        encodings = ['latin-1', 'cp1252', 'iso-8859-1', 'utf-8']
        rtf_text = None
        
        for encoding in encodings:
            try:
                candidate = rtf_content.decode(encoding, errors='ignore')
                parsed = rtf_to_text(candidate)
                if parsed.strip():
                    rtf_text = candidate
                    break
            except Exception:
                continue
        
        if rtf_text is None:
            rtf_text = rtf_content.decode('latin-1', errors='ignore')
        
        # Usa striprtf para extrair texto limpo
        try:
            plain_text = rtf_to_text(rtf_text)
        except Exception as e:
            # Fallback: extração manual básica removendo comandos RTF
            plain_text = rtf_text
            # Remove grupos RTF vazios primeiro
            while '{' in plain_text and '}' in plain_text:
                plain_text = re.sub(r'\{[^{}]*\}', '', plain_text)
            # Remove comandos RTF
            plain_text = re.sub(r'\\[a-z]+\d*\s*', ' ', plain_text)
            plain_text = re.sub(r'\\[{}]', '', plain_text)
            # Remove caracteres especiais RTF
            plain_text = re.sub(r'\\\'[0-9a-f]{2}', '', plain_text)
            # Remove números soltos que são comandos RTF
            plain_text = re.sub(r'\s+\d+\s+', ' ', plain_text)
        
        # Limpa o texto extraído
        # Remove linhas que são apenas nomes de fontes ou comandos
        lines = plain_text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            
            # Remove caracteres de controle
            line = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', line)
            line = line.strip()
            
            # Remove linhas que são apenas artefatos
            if not line:
                cleaned_lines.append("")
                continue
            
            # Remove linhas que são apenas nomes de fontes ou comandos RTF
            if (line.lower() in ['times new roman', 'arial', 'calibri', 'helvetica', 
                                 'trebuchet ms', 'cambria', 'times'] or
                re.match(r'^[a-z\s]+\}?$', line.lower()) or
                re.match(r'^[\d\s\-]+$', line) or
                line.count('}') > line.count(' ') or
                (len(line) < 3 and not line.isalnum())):
                continue
            
            # Remove artefatos de conversão RTF
            line = re.sub(r'\s+', ' ', line)  # Múltiplos espaços
            line = re.sub(r'^[a-z]+\s+[a-z]+\s+', '', line)  # Remove palavras soltas no início
            line = re.sub(r'\s*\}\s*$', '', line)  # Remove chaves no final
            line = re.sub(r'^\s*\{\s*', '', line)  # Remove chaves no início
            
            # Remove números soltos no início/fim
            line = re.sub(r'^\s*[\d\-]+\s+', '', line)
            line = re.sub(r'\s+[\d\-]+\s*$', '', line)
            
            if line.strip():
                cleaned_lines.append(line.strip())
        
        # Processa as linhas limpas e aplica formatação
        markdown_lines = []
        
        for line in cleaned_lines:
            if not line:
                markdown_lines.append("")
                continue
            
            # Detecta títulos principais (maiúsculas, sem ponto final, contém palavras-chave)
            if (line.isupper() and 15 < len(line) < 120 and 
                not line.endswith('.') and 
                ('TOMOGRAFIA' in line or 'ANGIO' in line or 'COMPUTADORIZADA' in line)):
                markdown_lines.append(f"## {line}")
            # Detecta seções importantes
            elif any(keyword in line.upper() for keyword in 
                    ['INDICAÇÃO CLÍNICA', 'TÉCNICA DO EXAME', 'ASPECTOS OBSERVADOS', 'IMPRESSÃO']):
                # Se está em maiúsculas, é título
                if line.isupper() and len(line) > 10:
                    markdown_lines.append(f"## {line}")
                else:
                    # Se começa com palavra-chave, aplica negrito
                    for keyword in ['INDICAÇÃO', 'TÉCNICA', 'ASPECTOS', 'IMPRESSÃO']:
                        if line.upper().startswith(keyword):
                            markdown_lines.append(f"**{line}**")
                            break
                    else:
                        markdown_lines.append(line)
            # Detecta texto em itálico (geralmente notas de rodapé)
            elif ('probabilidade' in line.lower() or 
                  'médico' in line.lower() or 
                  'diagnóstica' in line.lower()):
                markdown_lines.append(f"*{line}*")
            else:
                markdown_lines.append(line)
        
        # Limpa linhas vazias excessivas
        result = []
        prev_empty = False
        for line in markdown_lines:
            if not line.strip():
                if not prev_empty:
                    result.append("")
                prev_empty = True
            else:
                result.append(line)
                prev_empty = False
        
        return "\n".join(result)
    
    except Exception as e:
        import traceback
        return f"# Erro ao converter {rtf_path}\n\nErro: {str(e)}\n\nTraceback: {traceback.format_exc()}"


def main():
    """Função principal que processa todos os arquivos na pasta Reports."""
    reports_dir = Path(__file__).parent / "Reports"
    
    if not reports_dir.exists():
        print(f"Erro: Pasta {reports_dir} não encontrada!")
        return
    
    # Cria pasta para markdown
    markdown_dir = reports_dir.parent / "Reports_markdown"
    markdown_dir.mkdir(exist_ok=True)
    
    # Processa arquivos .docx
    docx_files = list(reports_dir.glob("*.docx"))
    print(f"Encontrados {len(docx_files)} arquivos .docx")
    
    for docx_file in docx_files:
        print(f"Convertendo {docx_file.name}...")
        markdown_content = convert_docx_to_markdown(docx_file)
        output_file = markdown_dir / f"{docx_file.stem}.md"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        print(f"  ✓ Salvo em {output_file.name}")
    
    # Processa arquivos .rtf
    rtf_files = list(reports_dir.glob("*.rtf"))
    print(f"\nEncontrados {len(rtf_files)} arquivos .rtf")
    
    for rtf_file in rtf_files:
        print(f"Convertendo {rtf_file.name}...")
        markdown_content = convert_rtf_to_markdown(rtf_file)
        output_file = markdown_dir / f"{rtf_file.stem}.md"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        print(f"  ✓ Salvo em {output_file.name}")
    
    print(f"\n✓ Conversão concluída! Arquivos salvos em {markdown_dir}")


if __name__ == "__main__":
    main()
